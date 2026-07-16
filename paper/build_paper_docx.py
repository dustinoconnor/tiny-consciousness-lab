#!/usr/bin/env python3
"""Build the audited research-paper draft as a deterministic DOCX."""

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "paper" / "RESEARCH_PAPER_DRAFT.md"
OUTPUT = ROOT / "paper" / "RESEARCH_PAPER_DRAFT.docx"

INK = RGBColor(26, 45, 58)
BLUE = RGBColor(38, 87, 112)
MUTED = RGBColor(92, 103, 110)
LIGHT = RGBColor(220, 228, 233)
BODY_FONT = "Calibri"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def create_list_numbering(doc, kind):
    """Create a fresh, properly spaced single-level list definition."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "number" else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if kind == "number" else "•")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "300")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_list_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)


def set_run(run, size=11, bold=None, italic=None, color=None, font=BODY_FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_rule(paragraph, color="B7C7D0", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr, fld_char2))


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 31, 31)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, INK),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "Abstract" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Abstract", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Abstract"]
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(42, 48, 52)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.left_indent = Inches(0.28)
    style.paragraph_format.right_indent = Inches(0.28)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.25


def add_inline_markdown(paragraph, text, size=11, color=None):
    token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            set_run(paragraph.add_run(text[cursor:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("*"):
            set_run(paragraph.add_run(token[1:-1]), size=size, italic=True, color=color)
        elif token.startswith("`"):
            set_run(paragraph.add_run(token[1:-1]), size=size - 0.5, color=color, font="Courier New")
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(label)
            set_run(run, size=size, color=BLUE)
            run.underline = True
            hyperlink = OxmlElement("w:hyperlink")
            rel_id = paragraph.part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink.set(qn("r:id"), rel_id)
            paragraph._p.remove(run._r)
            hyperlink.append(run._r)
            paragraph._p.append(hyperlink)
        cursor = match.end()
    if cursor < len(text):
        set_run(paragraph.add_run(text[cursor:]), size=size, color=color)


def add_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run(kicker.add_run("AUDIT-FIRST RESEARCH PAPER DRAFT"), size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run(title.add_run("Grounded Recurrent Control\nin Toy Partially Observable Worlds"), size=28, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    set_run(
        subtitle.add_run(
            "A thesis-driven synthesis of recurrence, grounded valence, attention, world models, "
            "hierarchical access, maintenance, and embodied predictive control"
        ),
        size=13,
        color=BLUE,
    )

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    set_paragraph_rule(rule)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    set_run(meta.add_run("Tiny Consciousness Lab"), size=12, bold=True, color=INK)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(date.add_run("16 July 2026  ·  Local repository audit"), size=10.5, color=MUTED)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(42)
    note.paragraph_format.left_indent = Inches(0.55)
    note.paragraph_format.right_indent = Inches(0.55)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        note.add_run(
            "Scope statement: This paper evaluates an engineered operational access-consciousness "
            "architecture in toy systems. It does not claim phenomenal consciousness, sentience, "
            "AGI, or biological equivalence."
        ),
        size=10.5,
        italic=True,
        color=MUTED,
    )
    doc.add_page_break()


def configure_page_furniture(section):
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("TINY CONSCIOUSNESS LAB  ·  RESEARCH PAPER DRAFT"), size=8.5, bold=True, color=MUTED)
    set_paragraph_rule(p, color="D5DEE3", size="4")

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_figure(doc, line):
    match = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", line)
    if not match:
        return
    caption, path = match.groups()
    image_path = Path(path)
    if not image_path.is_absolute():
        image_path = (SOURCE.parent / image_path).resolve()
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(6.25))
    doc_pr = inline._inline.docPr
    doc_pr.set("title", caption.split(".", 1)[0])
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    add_inline_markdown(cap, caption, size=9, color=MUTED)


def add_body_from_markdown(doc):
    lines = SOURCE.read_text().splitlines()
    # Skip title/subtitle/byline already represented on the cover.
    index = next(i for i, line in enumerate(lines) if line.strip() == "### Abstract")
    in_abstract = False
    buffer = []
    active_list_kind = None
    active_list_id = None

    def flush():
        nonlocal buffer
        if not buffer:
            return
        text = " ".join(part.strip() for part in buffer).strip()
        if text:
            style = "Abstract" if in_abstract else "Normal"
            p = doc.add_paragraph(style=style)
            add_inline_markdown(p, text, size=10.5 if in_abstract else 11)
        buffer = []

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush()
            active_list_kind = None
            active_list_id = None
            index += 1
            continue
        if line.startswith("!["):
            flush()
            active_list_kind = None
            add_figure(doc, line)
        elif line.startswith("### "):
            flush()
            active_list_kind = None
            text = line[4:]
            if text == "Abstract":
                in_abstract = True
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(7)
                p.paragraph_format.keep_with_next = True
                set_run(p.add_run("Abstract"), size=13, bold=True, color=BLUE)
            else:
                doc.add_paragraph(text, style="Heading 2")
        elif line.startswith("## "):
            flush()
            active_list_kind = None
            in_abstract = False
            doc.add_paragraph(line[3:], style="Heading 1")
        elif re.match(r"^\d+\. ", line):
            flush()
            if active_list_kind != "number":
                active_list_kind = "number"
                active_list_id = create_list_numbering(doc, "number")
            p = doc.add_paragraph(style="List Number")
            apply_list_numbering(p, active_list_id)
            add_inline_markdown(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            flush()
            if active_list_kind != "bullet":
                active_list_kind = "bullet"
                active_list_id = create_list_numbering(doc, "bullet")
            p = doc.add_paragraph(style="List Bullet")
            apply_list_numbering(p, active_list_id)
            add_inline_markdown(p, line[2:])
        elif line.startswith("> "):
            flush()
            active_list_kind = None
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.line_spacing = 1.25
            set_paragraph_rule(p, color="B7C7D0", size="8")
            add_inline_markdown(p, line[2:], size=10.5, color=INK)
        elif line.startswith("**Keywords:**"):
            flush()
            active_list_kind = None
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(14)
            add_inline_markdown(p, line, size=10)
            in_abstract = False
        else:
            active_list_kind = None
            buffer.append(line)
        index += 1
    flush()


def build():
    doc = Document()
    section = doc.sections[0]
    configure_page_furniture(section)
    configure_styles(doc)
    add_cover(doc)
    add_body_from_markdown(doc)
    props = doc.core_properties
    props.title = "Grounded Recurrent Control in Toy Partially Observable Worlds"
    props.subject = "Thesis-driven audit of a regulated Functional Ego architecture"
    props.author = "Tiny Consciousness Lab"
    props.keywords = "functional ego, access consciousness, recurrent control, grounded valence, hierarchical workspace, Unity, MPC"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
